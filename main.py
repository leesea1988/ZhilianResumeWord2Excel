# -*- coding: utf-8 -*-
"""
本脚本实现从简历文件夹的docx文件的相关数据
导出到根目录下excel表格
docx是智联招聘上导出的word文件

"""
import os
import sys


reload(sys)
sys.setdefaultencoding('utf8')  # 编译环境utf8
from glob import glob
import re
import time
import logging
logging.basicConfig(level=logging.DEBUG,#控制台打印的日志级别
                    filename='resumelist.log',
                    filemode='a',##模式，有w和a，w就是写模式，每次都会重新写日志，覆盖之前的日志
                    #a是追加模式，默认如果不写的话，就是追加模式
                    format=
                    '%(asctime)s - %(pathname)s[line:%(lineno)d] - %(levelname)s: %(message)s'
                    #日志格式
                    )


try:
    from docx import Document
except ImportError:
    logging.error(u'缺少模块python-docx，正在自动安装')
    import subprocess

    p = subprocess.Popen('easy_install python-docx', shell=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    logging.debug(p.stdout.readlines())
    for line in p.stdout.readlines():
        logging.debug(line)
    retval = p.wait()
    from docx import Document
    # raise
try:
    import openpyxl
except ImportError:
    logging.debug(u'缺少模块openpyxl，正在自动安装')
    import subprocess

    p = subprocess.Popen('easy_install openpyxl', shell=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    logging.debug(p.stdout.readlines())
    for line in p.stdout.readlines():
        logging.debug(line)
    retval = p.wait()
    import openpyxl
    # raise
##################################这是彩色打印
import ctypes

STD_INPUT_HANDLE = -10
STD_OUTPUT_HANDLE = -11
STD_ERROR_HANDLE = -12

FOREGROUND_BLACK = 0x0
FOREGROUND_BLUE = 0x01  # text color contains blue.
FOREGROUND_GREEN = 0x02  # text color contains green.
FOREGROUND_RED = 0x04  # text color contains red.
FOREGROUND_INTENSITY = 0x08  # text color is intensified.

BACKGROUND_BLUE = 0x10  # background color contains blue.
BACKGROUND_GREEN = 0x20  # background color contains green.
BACKGROUND_RED = 0x40  # background color contains red.
BACKGROUND_INTENSITY = 0x80  # background color is intensified.


# 上面这一大段都是在设置前景色和背景色，其实可以用数字直接设置，我的代码直接用数字设置颜色


class Color:
    std_out_handle = ctypes.windll.kernel32.GetStdHandle(STD_OUTPUT_HANDLE)

    def set_cmd_color(self, color, handle=std_out_handle):
        bool = ctypes.windll.kernel32.SetConsoleTextAttribute(handle, color)
        return bool

    def reset_color(self):
        self.set_cmd_color(FOREGROUND_RED | FOREGROUND_GREEN | FOREGROUND_BLUE | FOREGROUND_INTENSITY)
        # 初始化颜色为黑色背景，纯白色字，CMD默认是灰色字体的

    def print_red_text(self, print_text):
        self.set_cmd_color(4 | 8)
        logging.debug(print_text)
        self.reset_color()
        # 红色字体

    def print_green_text(self, print_text):
        self.set_cmd_color(FOREGROUND_GREEN | FOREGROUND_INTENSITY)
        # c = raw_input(print_text.encode('gbk'))
        # c = raw_input(print_text)
        logging.debug(print_text)
        self.reset_color()
        # return c

    def print_yellow_text(self, print_text):
        self.set_cmd_color(6 | 8)
        logging.debug(print_text)
        self.reset_color()
        # 黄色字体

    def print_blue_text(self, print_text):
        self.set_cmd_color(1 | 10)
        logging.debug(print_text)
        self.reset_color()
        # 蓝色字体


clr = Color()
clr.set_cmd_color(FOREGROUND_RED | FOREGROUND_GREEN | FOREGROUND_BLUE | FOREGROUND_INTENSITY)
# clr.print_red_text('red')
# clr.print_green_text("green")
# clr.print_blue_text('blue')
# clr.print_yellow_text('yellow')
##########################################


PROJECT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))

#所有要抽取内容中，可以匹配到的正则表达式
re_pats = {

    'sex': u'^男|女',
    'age': u'\d+岁',
    'workyear': u'\d+年工作经验',
    'education': u'^本科|硕士|博士',
    'university': u'大学$',
    'major': u'^男',
    'experience': u'^男',
    'birth': u'\d+年\d+月',
}


def get_filename_by_path(path, forbid_word=''):
    searched_filenames = glob(path)
    return [i for i in searched_filenames if forbid_word not in i]


def parse_text_by_repat(text, re_pat):
    find = re.findall(re_pat, text)
    return (find[0].strip() if find else None)


def read_docx(filename):
    result = {
        'name': None,
        'sex': None,
        'age': None,
        'workyear': None,
        'education': None,
        'university': None,
        'major': None,
        'experience': None,
        'birth': None
    }

    # 简历的标题内容，包括名字、应聘职位、工作地点等
    title = {
        'name': None,
    }

    # 基本信息内容，包括性别、年龄、工作经验等
    basic_info = {
        'sex': None,
        'age': None,
        'workyear': None,
        'birth': None,
        'education': None,

    }

    # 教育经历，包括毕业学校、专业
    educate = {
        'university': None,
        'major': None,
        'education': None,
    }

    #工作经验
    experience = {
        'experience':None,
    }

    document = Document(filename)

    # 读取标题信息 ，名字一般固定在段落5，这里写死
    employee_name = document.paragraphs[5]
    title['name'] = employee_name.text
    result['name'] = title['name']

    #读取基本信息
    locate_sex_male = u'男｜'  # 定位基本信息段落的关键字
    locate_sex_female = u'女｜'  # 定位基本信息段落的关键字

    max_search_para = 20  # 最多搜索前n段来位段落
    logging.debug(u'定位并解析数据')
    for para_num, para in enumerate(document.paragraphs):
        if para_num < max_search_para and (locate_sex_male in para.text or locate_sex_female in para.text ):

            for key in basic_info.keys():
                basic_info[key] = parse_text_by_repat(para.text, re_pats[key])

            #将基本信息内容复制到返回数组中
            for key in basic_info.keys():
                result[key] = basic_info[key]

            # 包含关键字则退出循环搜索
            break
        elif para_num >= max_search_para:
            logging.error(u'当前应聘者基本信息读取失败，没有性别信息！')
            raise IOError(u'Error！应聘者基本信息读取失败--文件前%s段不包含指定的文字：%s，请检查！' % (max_search_para, locate_sex_male))

    #遍历所有段落  读取教育经历
    for para_num , para in enumerate(document.paragraphs):

        if  u'教育经历' == para.text:
            num = para_num
            university_major =  document.paragraphs[num + 2]
            educate['university'] = university_major.text.split("   ")[0]
            educate['major'] = university_major.text.split("   ")[1]
            education = document.paragraphs[num + 3]
            educate['education'] = education.text.split("    ")[1]

            for key in educate.keys():
                result[key] = educate[key]

            #包含关键字则退出循环搜索
            break

    # 获取第一个表格  读取工作经验
    if(document.tables):
        experience_tab = document.tables[0]
        str_experience = ""
        for row_number in range(0,len(experience_tab.rows)):
            if(experience_tab.cell(row_number,0).text!=u'' and not experience_tab.cell(row_number,0).text.startswith(u'工作描述')):
                str_experience = str_experience +experience_tab.cell(row_number,0).text + u'\n'
        result['experience'] = str_experience;
    return result

def write_excel(excel_name, result_dicts):
    from openpyxl.workbook import Workbook

    # ExcelWriter,里面封装好了对Excel的写操作
    from openpyxl.writer.excel import ExcelWriter

    # get_column_letter函数将数字转换为相应的字母，如1-->A,2-->B
    from openpyxl.cell import get_column_letter

    from openpyxl.reader.excel import load_workbook

    if os.path.isfile(excel_name):
        # #读取excel2007文件
        wb = load_workbook(excel_name)
    else:
        # 新建一个workbook
        wb = Workbook()

    # 新建一个excelWriter
    ew = ExcelWriter(workbook=wb)

    # 设置文件输出路径与名称
    dest_filename = excel_name

    # # 获取第一个sheet
    try:
        ws = wb.get_sheet_by_name('sheet1')
    except KeyError:
        ws = wb.worksheets[0]
        ws.title = "sheet1"

    # 第一个sheet是ws
    # ws = wb.worksheets[0]

    # #设置ws的名称
    # ws.title = "sheet1"

    line = 1
    logging.debug(u'定位写入坐标')
    while ws.cell("A%s" % line).value:
         line += 1
    logging.debug(u'从第%s行开始写入' % line)

    if not os.path.isfile(excel_name):
        ws.cell("A%s" % line).value = u'序号'
        ws.cell("B%s" % line).value = u'姓名'
        ws.cell("C%s" % line).value = u'性别'
        ws.cell("D%s" % line).value = u'年龄'
        ws.cell("E%s" % line).value = u'工作年限'
        ws.cell("F%s" % line).value = u'学历'
        ws.cell("G%s" % line).value = u'毕业院校'
        ws.cell("H%s" % line).value = u'专业'
        ws.cell("I%s" % line).value = u'工作经历'

        line += 1
    for i, result in enumerate(result_dicts):
        logging.debug(u'正在写入第%s条数据到excel' % (i + 1))
        ws.cell("A%s" % line).value = line-1
        ws.cell("B%s" % line).value = result['name']
        ws.cell("C%s" % line).value = result['sex']
        ws.cell("D%s" % line).value = result['age']
        ws.cell("E%s" % line).value = result['workyear']
        ws.cell("F%s" % line).value = result['education']
        ws.cell("G%s" % line).value = result['university']
        ws.cell("H%s" % line).value = result['major']
        ws.cell("I%s" % line).value = result['experience']

        line += 1

    # 最后保存文件
    ew.save(filename=excel_name)


def main():
    logging.debug(u'开始执行')
    logging.debug(u'从"简历"文件夹查找docx文件')
    filenames = get_filename_by_path(u'简历/*.docx', '~$')
    result_dicts = []
    for filename in filenames:
        logging.debug(u'读取文件：')
        clr.print_blue_text(os.path.basename(filename))
        result_dicts.append(read_docx(filename))
    save_filename = 'ResumeList.xlsx'
    write_excel(save_filename, result_dicts)
    logging.debug(u'执行完毕，文件保存至')
    clr.print_blue_text(save_filename)


if __name__ == '__main__':
    main()
